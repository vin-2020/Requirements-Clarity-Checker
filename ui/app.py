# ui/app.py
import streamlit as st
import sys
import os
import re
import docx
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import pandas as pd
import inspect
import importlib
import sqlite3  # <-- to catch IntegrityError

# Make local packages importable when run from /ui
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# ------------------------- Imports from your app -------------------------
# Analyzer functions (some may not exist in older versions; we shim below)
from core.analyzer import (
    check_requirement_ambiguity,
    check_passive_voice,
    check_incompleteness,
)
try:
    from core.analyzer import check_singularity  # optional
except Exception:
    def check_singularity(_text: str):
        return []  # safe fallback

# Scoring (signature may be old or new; we handle both)
from core.scoring import calculate_clarity_score

# Optional rule engine (new); fall back to a dummy if missing
try:
    from core.rule_engine import RuleEngine
except Exception:
    class RuleEngine:
        """Minimal stub to keep the app running if core.rule_engine is absent."""
        def __init__(self):
            pass

# LLM helpers (robust, reloadable)
from llm import ai_suggestions as ai  # module import
ai = importlib.reload(ai)  # ensure latest functions are picked up

# Core helpers (these should exist)
get_ai_suggestion = getattr(ai, "get_ai_suggestion")
generate_requirement_from_need = getattr(ai, "generate_requirement_from_need")

# Optional helpers with safe fallbacks
get_chatbot_response = getattr(
    ai,
    "get_chatbot_response",
    # Fallback: reuse get_ai_suggestion with a stitched history
    lambda api_key, history: get_ai_suggestion(
        api_key,
        "\n".join(
            f"{m.get('role','user').upper()}: {(m.get('parts') or [m.get('content','')])[0]}"
            for m in history
        ) + "\nASSISTANT:"
    ),
)

# New decomposition helper
decompose_requirement_with_ai = getattr(
    ai, "decompose_requirement_with_ai",
    lambda api_key, requirement_text: "Decomposition helper failed to load."
)

# --- Safe import for AI extractor ---
try:
    from llm.ai_suggestions import extract_requirements_with_ai
    HAS_AI_PARSER = True
except Exception:
    HAS_AI_PARSER = False
    def extract_requirements_with_ai(*args, **kwargs):
        return []  # fallback

# Database helpers  (DB memory integration)
from db.database import init_db, add_project, get_all_projects  # type: ignore
from db import database as db  # type: ignore  # <-- module import avoids name errors
db = importlib.reload(db)  # ensure latest functions (add_document, etc.) are present

# ========================= Helpers for Analyzer =========================

def _read_docx_text_and_rows(uploaded_file):
    """
    Returns (flat_text, table_rows) where:
      - flat_text is all paragraph + cell text joined with newlines
      - table_rows is a list of rows (each row = list[str] of cell texts)
    """
    d = docx.Document(uploaded_file)
    parts = []

    # paragraphs
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    # tables
    rows = []
    for tbl in d.tables:
        for r in tbl.rows:
            row_cells = []
            for c in r.cells:
                cell_text = " ".join(p.text.strip() for p in c.paragraphs if p.text.strip())
                row_cells.append(cell_text)
                if cell_text:
                    parts.append(cell_text)
            rows.append(row_cells)

    flat_text = "\n".join(parts)
    return flat_text, rows

# --- read .docx from disk path (for re-analysis of stored files) ---
def _read_docx_text_and_rows_from_path(path: str):
    d = docx.Document(path)
    parts = []
    rows = []

    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    for tbl in d.tables:
        for r in tbl.rows:
            row_cells = []
            for c in r.cells:
                cell_text = " ".join(p.text.strip() for p in c.paragraphs if p.text.strip())
                row_cells.append(cell_text)
                if cell_text:
                    parts.append(cell_text)
            rows.append(row_cells)

    flat_text = "\n".join(parts)
    return flat_text, rows

def _extract_requirements_from_table_rows(table_rows):
    """
    If the DOCX has a requirements table with headers like:
      ID | Requirement | Rationale | Verification | Acceptance Criteria
    this will harvest (id, text) pairs robustly.
    """
    if not table_rows:
        return []

    def _norm(s): return (s or "").strip().lower()

    # Find a header row
    header_idx = None
    for i, row in enumerate(table_rows):
        cells = [_norm(c) for c in row]
        if not cells:
            continue
        if ("id" in cells[0] and any("requirement" in c for c in cells)):
            header_idx = i
            break
        if ("requirement" in cells[0] and any(c == "id" for c in cells)):
            header_idx = i
            break

    if header_idx is None:
        return []

    # Map columns
    header = [_norm(c) for c in table_rows[header_idx]]
    id_col = None
    req_col = None
    for idx, h in enumerate(header):
        if h == "id":
            id_col = idx
        if "requirement" in h:
            req_col = idx
    if id_col is None or req_col is None:
        return []

    # IDs like SAT-REQ-001, ABC-123, SUBSYS-THERM-42
    id_pat = re.compile(r'^[A-Z][A-Z0-9-]*-\d+$')

    out = []
    for row in table_rows[header_idx + 1:]:
        if len(row) <= max(id_col, req_col):
            continue
        rid = (row[id_col] or "").strip()
        rtx = (row[req_col] or "").strip()
        if not rid or not rtx:
            continue
        if id_pat.match(rid):
            out.append((rid, rtx))
    return out

def extract_requirements_from_string(content: str):
    """Extract (id, text) pairs like 'SAT-REQ-001 ...' or '1.' lines."""
    requirements = []
    req_pattern = re.compile(r'^(([A-Z][A-Z0-9-]*-\d+)|(\d+\.))\s+(.*)$')
    for line in content.split('\n'):
        line = line.strip()
        m = req_pattern.match(line)
        if m:
            rid = m.group(1)
            text = m.group(4)
            requirements.append((rid, text))
    return requirements

def extract_requirements_from_file(uploaded_file):
    """Read .txt/.docx to text, then parse requirements."""
    if uploaded_file.name.endswith('.txt'):
        content = uploaded_file.getvalue().decode("utf-8")
        table_rows = []
    elif uploaded_file.name.endswith('.docx'):
        content, table_rows = _read_docx_text_and_rows(uploaded_file)
    else:
        content, table_rows = "", []

    reqs = _extract_requirements_from_table_rows(table_rows)
    if reqs:
        return reqs
    return extract_requirements_from_string(content)

def format_requirement_with_highlights(req_id, req_text, issues):
    """Inline HTML highlight for ambiguous/passive elements."""
    highlighted_text = req_text
    if issues.get('ambiguous'):
        for word in issues['ambiguous']:
            highlighted_text = re.sub(
                r'\b' + re.escape(word) + r'\b',
                f'<span style="background-color:#FFFF00;color:black;padding:2px 4px;border-radius:3px;">{word}</span>',
                highlighted_text,
                flags=re.IGNORECASE
            )
    if issues.get('passive'):
        for phrase in issues['passive']:
            highlighted_text = re.sub(
                re.escape(phrase),
                f'<span style="background-color:#FFA500;padding:2px 4px;border-radius:3px;">{phrase}</span>',
                highlighted_text,
                flags=re.IGNORECASE
            )

    display_html = f"⚠️ <strong>{req_id}</strong> {highlighted_text}"
    explanations = []
    if issues.get('ambiguous'):
        explanations.append(f"<i>- Ambiguity: Found weak words: <b>{', '.join(issues['ambiguous'])}</b></i>")
    if issues.get('passive'):
        explanations.append(f"<i>- Passive Voice: Found phrase: <b>'{', '.join(issues['passive'])}'</b>. Consider active voice.</i>")
    if issues.get('incomplete'):
        explanations.append("<i>- Incompleteness: Requirement appears to be a fragment.</i>")
    if issues.get('singularity'):
        explanations.append(f"<i>- Singularity: Multiple actions indicated: <b>{', '.join(issues['singularity'])}</b></i>")
    if explanations:
        display_html += "<br>" + "<br>".join(explanations)

    return (
        f'<div style="background-color:#FFF3CD;color:#856404;padding:10px;'
        f'border-radius:5px;margin-bottom:10px;">{display_html}</div>'
    )

def safe_call_ambiguity(text: str, engine: RuleEngine | None):
    """Call check_requirement_ambiguity with or without rule engine, depending on its signature."""
    try:
        return check_requirement_ambiguity(text, engine)
    except TypeError:
        return check_requirement_ambiguity(text)

def safe_clarity_score(total_reqs: int, results: list[dict], issue_counts=None, engine: RuleEngine | None = None):
    """
    Call calculate_clarity_score supporting both signatures:
    - New: calculate_clarity_score(total_reqs, issue_counts, rule_engine)
    - Old: calculate_clarity_score(total_reqs, flagged_reqs)
    """
    try:
        sig = inspect.signature(calculate_clarity_score)
        if len(sig.parameters) >= 3:
            return calculate_clarity_score(total_reqs, issue_counts or {}, engine)
        else:
            flagged_reqs = sum(1 for r in results if r['ambiguous'] or r['passive'] or r['incomplete'])
            return calculate_clarity_score(total_reqs, flagged_reqs)
    except Exception:
        flagged_reqs = sum(1 for r in results if r['ambiguous'] or r['passive'] or r['incomplete'])
        clear_reqs = max(0, total_reqs - flagged_reqs)
        return int((clear_reqs / total_reqs) * 100) if total_reqs else 100

def _open_db_conn():
    """
    Open a connection to the SAME SQLite DB the db module is using.
    Tries, in order:
      - db.get_connection() (reuse existing connection)
      - db.DB_PATH / DB_FILE / DB_NAME / DATABASE_PATH (string path)
      - common filenames next to db module: reqcheck.db, database.db, app.db
    """
    if hasattr(db, "get_connection"):
        try:
            conn = db.get_connection()
            conn.execute("PRAGMA foreign_keys = ON;")
            return conn, False
        except Exception:
            pass

    for attr in ("DB_PATH", "DB_FILE", "DB_NAME", "DATABASE_PATH"):
        path = getattr(db, attr, None)
        if isinstance(path, str) and path:
            conn = sqlite3.connect(path)
            conn.execute("PRAGMA foreign_keys = ON;")
            return conn, True

    base_dir = os.path.dirname(db.__file__)
    for name in ("reqcheck.db", "database.db", "app.db"):
        candidate = os.path.join(base_dir, name)
        if os.path.exists(candidate):
            conn = sqlite3.connect(candidate)
            conn.execute("PRAGMA foreign_keys = ON;")
            return conn, True

    raise RuntimeError(
        "Could not locate the SQLite DB file. "
        "Expose DB_PATH in db.database or provide db.get_connection()."
    )

# --- File storage helpers ---
def _sanitize_filename(name: str) -> str:
    return re.sub(r'[^A-Za-z0-9._-]+', '_', name)

def _save_uploaded_file_for_doc(project_id: int, doc_id: int, display_name: str, uploaded_file) -> str:
    base_dir = os.path.join("data", "projects", str(project_id), "documents")
    os.makedirs(base_dir, exist_ok=True)
    safe_name = _sanitize_filename(display_name)
    out_path = os.path.join(base_dir, f"{doc_id}_{safe_name}")
    try:
        uploaded_file.seek(0)
        with open(out_path, "wb") as f:
            f.write(uploaded_file.read())
    except Exception:
        with open(out_path, "wb") as f:
            f.write(uploaded_file.getvalue())
    return out_path

# =============================== UI Setup ===============================
st.set_page_config(
    page_title="ReqCheck Workspace",
    page_icon="https://github.com/vin-2020/Requirements-Clarity-Checker/blob/main/Logo.png?raw=true",
    layout="wide"
)

# Global CSS
st.markdown("""
<style>
    .req-container { padding:10px;border-radius:5px;margin-bottom:10px;border:1px solid #ddd; }
    .flagged { background:#FFF3CD;color:#856404;border-color:#FFEEBA; }
    .clear { background:#D4EDDA;color:#155724;border-color:#C3E6CB; }
    .highlight-ambiguity { background:#FFFF00;color:black;padding:2px 4px;border-radius:3px; }
    .highlight-passive { background:#FFA500;padding:2px 4px;border-radius:3px; }
    .explanation { font-size:0.9em;font-style:italic;color:#6c757d;margin-top:5px; }
</style>
""", unsafe_allow_html=True)

with st.sidebar:
    st.image("https://github.com/vin-2020/Requirements-Clarity-Checker/blob/main/ReqCheck_Logo.png?raw=true", use_container_width=True)
    st.header("About ReqCheck")
    st.info("An AI-assisted tool to evaluate the quality of system requirements...")
    st.header("Project Links")
    st.markdown("[GitHub Repository](https://github.com/vin-2020/Requirements-Clarity-Checker)")
    st.markdown("[INCOSE Handbook](https://www.incose.org/products-and-publications/se-handbook)")

st.title("✨ ReqCheck: AI-Powered Requirements Assistant")

# ✅ Initialize the database on first run (DB memory)
init_db()

# Single API key stored globally
if 'api_key' not in st.session_state:
    st.session_state.api_key = ''
api_key_input = st.text_input(
    "Enter your Google AI API Key to enable AI features:",
    type="password",
    value=st.session_state.api_key,
    key="api_key_global"
)
if api_key_input:
    st.session_state.api_key = api_key_input

# Track selected project globally
if 'selected_project' not in st.session_state:
    st.session_state.selected_project = None
st.markdown("Get your free API key from [Google AI Studio](https://aistudio.google.com/).")

# One RuleEngine instance (real or stub)
rule_engine = RuleEngine()

# ======================= Layout: main + right panel =======================
main_col, right_col = st.columns([4, 1], gap="large")

# ----------------------------- Right Panel (Projects) -----------------------------
with right_col:
    st.subheader("🗂️ Projects")

    if st.session_state.selected_project is not None:
        _pid, _pname = st.session_state.selected_project
        st.caption(f"Current: **{_pname}**")
        if st.button("Clear selection", key="btn_clear_proj_right"):
            st.session_state.selected_project = None
            st.rerun()

    projects = get_all_projects()
    names = [p[1] for p in projects] if projects else []
    if names:
        sel_name = st.selectbox("Open project:", names, key="proj_select_right")

        # Confirmation state
        if "confirm_delete" not in st.session_state:
            st.session_state.confirm_delete = False
        if "delete_project_id" not in st.session_state:
            st.session_state.delete_project_id = None
        if "delete_project_name" not in st.session_state:
            st.session_state.delete_project_name = None

        col1, col2 = st.columns(2)
        with col1:
            if st.button("Load", key="btn_load_proj_right"):
                for p in projects:
                    if p[1] == sel_name:
                        st.session_state.selected_project = p
                        st.success(f"Loaded: {sel_name}")
                        st.rerun()

        with col2:
            if st.button("Delete", key="btn_delete_proj_right"):
                for p in projects:
                    if p[1] == sel_name:
                        st.session_state.delete_project_id = p[0]
                        st.session_state.delete_project_name = sel_name
                        st.session_state.confirm_delete = True

        if st.session_state.confirm_delete:
            st.warning(
                f"You're about to delete '{st.session_state.delete_project_name}'. "
                "This cannot be undone."
            )
            c1, c2 = st.columns(2)
            with c1:
                if st.button("Confirm Delete", key="btn_confirm_delete_proj_right"):
                    pid_to_delete = st.session_state.delete_project_id

                    def _get_tables(conn):
                        cur = conn.cursor()
                        cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'")
                        return [r[0] for r in cur.fetchall()]

                    def _fk_refs(conn, ref_table_name):
                        """Return list of (table, from_col) pairs for all tables that have a FK to ref_table_name."""
                        refs = []
                        for t in _get_tables(conn):
                            try:
                                cur = conn.cursor()
                                cur.execute(f"PRAGMA foreign_key_list('{t}')")
                                for (id_, seq, table, from_col, to_col, on_update, on_delete, match) in cur.fetchall():
                                    if table == ref_table_name:
                                        refs.append((t, from_col))
                            except Exception:
                                pass
                        return refs

                    try:
                        conn, _should_close = _open_db_conn()
                        cur = conn.cursor()
                        cur.execute("PRAGMA foreign_keys = ON;")

                        # 1) Gather document ids for the project
                        cur.execute("SELECT id FROM documents WHERE project_id = ?", (pid_to_delete,))
                        doc_ids = [r[0] for r in cur.fetchall()]

                        if doc_ids:
                            # 2) Delete ALL rows in ANY table that references 'documents'
                            refs_to_documents = _fk_refs(conn, "documents")
                            for (tbl, col) in refs_to_documents:
                                qmarks = ",".join("?" for _ in doc_ids)
                                cur.execute(f"DELETE FROM {tbl} WHERE {col} IN ({qmarks})", doc_ids)

                            # 3) Now delete requirements explicitly, then documents
                            qmarks = ",".join("?" for _ in doc_ids)
                            try:
                                cur.execute(f"DELETE FROM requirements WHERE document_id IN ({qmarks})", doc_ids)
                            except Exception:
                                pass
                            cur.execute(f"DELETE FROM documents WHERE id IN ({qmarks})", doc_ids)

                        # 4) Delete ALL rows in ANY table that references 'projects'
                        refs_to_projects = _fk_refs(conn, "projects")
                        for (tbl, col) in refs_to_projects:
                            cur.execute(f"DELETE FROM {tbl} WHERE {col} = ?", (pid_to_delete,))

                        # 5) Finally delete the project
                        cur.execute("DELETE FROM projects WHERE id = ?", (pid_to_delete,))
                        conn.commit()
                        if _should_close:
                            conn.close()

                        st.success("Project deleted.")

                    except sqlite3.IntegrityError as e:
                        try:
                            cur.execute("PRAGMA foreign_key_check;")
                            problems = cur.fetchall()
                        except Exception:
                            problems = []
                        if problems:
                            st.error(f"Delete failed due to FK constraints. Offending rows: {problems}")
                        else:
                            st.error(f"Delete failed due to foreign key constraints: {e}")
                    except Exception as e:
                        st.error(f"Delete failed: {e}")

                    # clear state regardless
                    st.session_state.confirm_delete = False
                    st.session_state.delete_project_id = None
                    st.session_state.delete_project_name = None
                    st.rerun()

            with c2:
                if st.button("Cancel", key="btn_cancel_delete_proj_right"):
                    st.session_state.confirm_delete = False
                    st.session_state.delete_project_id = None
                    st.session_state.delete_project_name = None

    else:
        st.caption("No projects yet.")

    # Create new
    st.text_input("New project name:", key="new_proj_name_right")
    if st.button("Create", key="btn_create_proj_right"):
        new_name = st.session_state.get("new_proj_name_right", "").strip()
        if new_name:
            feedback = add_project(new_name)
            st.success(feedback)
            st.rerun()
        else:
            st.error("Please enter a project name.")

# ------------------------------ Main Tabs (LEAN) ------------------------------
# Import split tabs
from ui.tabs.analyzer_tab import render as render_analyzer_tab
from ui.tabs.needs_tab   import render as render_need_tab   # <-- plural
from ui.tabs.chat_tab    import render as render_chat_tab


with main_col:
    tab_analyze, tab_need, tab_chat = st.tabs([
        "📄 Document Analyzer",
        "💡 Need-to-Requirement Helper",
        "💬 Requirements Chatbot",
    ])

# Build a lightweight context so tab files can reuse your helpers without duplication.
CTX = {
    # flags
    "HAS_AI_PARSER": HAS_AI_PARSER,
    # AI helpers
    "get_ai_suggestion": get_ai_suggestion,
    "get_chatbot_response": get_chatbot_response,   # <-- make chatbot available to the tab
    "decompose_requirement_with_ai": decompose_requirement_with_ai,
    "extract_requirements_with_ai": extract_requirements_with_ai,
    # parsing/extract
    "_read_docx_text_and_rows": _read_docx_text_and_rows,
    "_read_docx_text_and_rows_from_path": _read_docx_text_and_rows_from_path,
    "_extract_requirements_from_table_rows": _extract_requirements_from_table_rows,
    "extract_requirements_from_string": extract_requirements_from_string,
    "extract_requirements_from_file": extract_requirements_from_file,
    # UI formatting
    "format_requirement_with_highlights": format_requirement_with_highlights,
    # analyzers
    "safe_call_ambiguity": safe_call_ambiguity,
    "check_passive_voice": check_passive_voice,
    "check_incompleteness": check_incompleteness,
    "check_singularity": check_singularity,
    "safe_clarity_score": safe_clarity_score,
    # storage
    "_save_uploaded_file_for_doc": _save_uploaded_file_for_doc,
    "_sanitize_filename": _sanitize_filename,
}

with tab_analyze:
    render_analyzer_tab(st, db, rule_engine, CTX)

with tab_need:
    render_need_tab(st, db, rule_engine, CTX)

with tab_chat:
    render_chat_tab(st, db, rule_engine, CTX)
